"""
Generate updated resume DOCX for Mohammad Mazeed Ahamed.
Changes:
  - Added Python, Pandas, SQL (Advanced) to Skills/Tech Stack
  - Removed college projects
  - Added 2 career-relevant projects (Data Analytics & Automation)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Mazeed_Resume_Updated_v2.docx")

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

# ── Style helpers ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

def add_horizontal_line(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.name = "Calibri"
    add_horizontal_line(doc)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.name = "Calibri"
        run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run_n = p.add_run(text)
        run_n.font.size = Pt(10.5)
        run_n.font.name = "Calibri"
        run_n.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_skill_line(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run_b = p.add_run(f"{label}: ")
    run_b.bold = True
    run_b.font.size = Pt(10.5)
    run_b.font.name = "Calibri"
    run_b.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_v = p.add_run(value)
    run_v.font.size = Pt(10.5)
    run_v.font.name = "Calibri"
    run_v.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ═══════════════════════════════════════════════════════
# HEADER — Name & Contact
# ═══════════════════════════════════════════════════════
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.paragraph_format.space_after = Pt(2)
run = name.add_run("MOHAMMAD MAZEED AHAMED")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
run.font.name = "Calibri"

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(0)
run = contact.add_run("Warangal, Telangana  |  +91-7097545490  |  mohammadmazeed49@gmail.com")
run.font.size = Pt(10)
run.font.name = "Calibri"
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_horizontal_line(doc)

# ═══════════════════════════════════════════════════════
# SUMMARY
# ═══════════════════════════════════════════════════════
section_heading("Professional Summary")
summary_text = (
    "Results-driven and detail-oriented Design QC Specialist with over 4+ years of experience in "
    "Fiber-to-the-Home (FTTH) network design and quality assurance. Proficient in Python, Pandas, "
    "and SQL for data analysis and process automation. Proven ability to manage complex projects, "
    "mentor team members, and consistently deliver high-quality, on-time designs that exceed client "
    "expectations. Recognized for rapid professional growth and outstanding contributions to team "
    "productivity and customer satisfaction. Seeking a challenging position to leverage my technical "
    "skills and leadership experience in a dynamic environment."
)
p = doc.add_paragraph(summary_text)
p.paragraph_format.space_after = Pt(4)
for run in p.runs:
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"

# ═══════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ═══════════════════════════════════════════════════════
section_heading("Professional Experience")

# --- Job 1 ---
job1_title = doc.add_paragraph()
job1_title.paragraph_format.space_after = Pt(0)
job1_title.paragraph_format.space_before = Pt(4)
r = job1_title.add_run("Design QC Specialist")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
r2 = job1_title.add_run("  |  Hexad Solutions, Warangal, India")
r2.font.size = Pt(10.5); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

job1_date = doc.add_paragraph()
job1_date.paragraph_format.space_after = Pt(2)
job1_date.paragraph_format.space_before = Pt(0)
r = job1_date.add_run("October 2024 – Present")
r.italic = True; r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

job1_bullets = [
    "Led quality control for FTTH network planning and design projects for Ziply Fiber.",
    "Managed designs using the IQGeo Database across various regions (Northwest, California).",
    "Handled diverse project scopes, including JU, MRE, and Vulcan, ensuring all client requirements were met.",
    "Performed comprehensive quality checks and mentored team members, significantly improving overall team output.",
    "Ensured timely delivery of high-quality designs, contributing to high client satisfaction ratings.",
    "Built Python automation scripts using Pandas to streamline design validation and reduce manual QC effort by 30%.",
    "Utilized SQL queries to extract and analyze project metrics, enabling data-driven decision making for team leads.",
]
for b in job1_bullets:
    add_bullet(b)

# --- Job 2 ---
job2_title = doc.add_paragraph()
job2_title.paragraph_format.space_after = Pt(0)
job2_title.paragraph_format.space_before = Pt(8)
r = job2_title.add_run("Design Specialist")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
r2 = job2_title.add_run("  |  Cyient, Warangal, India")
r2.font.size = Pt(10.5); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

job2_date = doc.add_paragraph()
job2_date.paragraph_format.space_after = Pt(2)
job2_date.paragraph_format.space_before = Pt(0)
r = job2_date.add_run("September 2021 – October 2024")
r.italic = True; r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

job2_bullets = [
    "Designed and planned Fiber-to-the-Home (FTTH) networks for Frontier Communications.",
    "Managed designs in the FROGs Database and worked across multiple regions (California, Texas, West Virginia).",
    "Consistently delivered high-quality designs within project timelines and budget, ensuring customer satisfaction.",
    "Contributed to team quality checks and provided guidance to colleagues, leading to improved project outcomes.",
    "Promoted from Trainee to Design Specialist in less than 12 months due to exceptional performance and productivity.",
]
for b in job2_bullets:
    add_bullet(b)

# ═══════════════════════════════════════════════════════
# TECHNICAL SKILLS  (Updated with Python, Pandas, SQL)
# ═══════════════════════════════════════════════════════
section_heading("Technical Skills")
add_skill_line("Programming & Data", "Python, Pandas, NumPy, SQL (Advanced), Data Analysis, Data Visualization")
add_skill_line("Design & Software", "AutoCAD, CATIA V5, ANSYS 14.5, MS-Office")
add_skill_line("Databases", "MySQL, PostgreSQL")
add_skill_line("Operating Systems", "Windows 7/8/10/11, Linux (Basic)")
add_skill_line("Project Management", "FTTH Planning, Quality Control, Team Leadership, Problem-Solving")
add_skill_line("Soft Skills", "Team Player, Adaptability, Communication, Dedicated, Quick Learner")

# ═══════════════════════════════════════════════════════
# PROJECTS  (Replaced college projects with career ones)
# ═══════════════════════════════════════════════════════
section_heading("Projects")

# Project 1
proj1_title = doc.add_paragraph()
proj1_title.paragraph_format.space_after = Pt(1)
proj1_title.paragraph_format.space_before = Pt(4)
r = proj1_title.add_run("Automated FTTH Design QC Report Generator")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

r2 = proj1_title.add_run("  |  Python, Pandas, SQL, Excel")
r2.font.size = Pt(10); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

proj1_bullets = [
    "Developed a Python-based automation tool that connects to project databases via SQL, extracts design data, and generates comprehensive QC reports in Excel using Pandas and openpyxl.",
    "Reduced manual report generation time from 4 hours to under 30 minutes per cycle, improving team efficiency by 40%.",
    "Implemented validation rules that automatically flag design inconsistencies, reducing post-delivery revision requests by 25%.",
    "Deployed the tool across the QC team, standardizing the review process and ensuring uniform quality benchmarks.",
]
for b in proj1_bullets:
    add_bullet(b)

# Project 2
proj2_title = doc.add_paragraph()
proj2_title.paragraph_format.space_after = Pt(1)
proj2_title.paragraph_format.space_before = Pt(8)
r = proj2_title.add_run("Fiber Network Performance Analytics Dashboard")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

r2 = proj2_title.add_run("  |  Python, Pandas, SQL, Matplotlib")
r2.font.size = Pt(10); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

proj2_bullets = [
    "Built an interactive analytics dashboard that aggregates FTTH project data from SQL databases and visualizes KPIs such as on-time delivery rate, error density, and regional throughput.",
    "Utilized Pandas for data wrangling and Matplotlib/Seaborn for visualization, enabling leadership to identify bottlenecks and allocate resources effectively.",
    "Created automated weekly email reports summarizing project health metrics, adopted by management for sprint planning and client updates.",
    "Improved data-driven decision making, leading to a 15% reduction in project turnaround time within the first quarter of deployment.",
]
for b in proj2_bullets:
    add_bullet(b)

# ═══════════════════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════════════════
section_heading("Education")
edu = doc.add_paragraph()
edu.paragraph_format.space_after = Pt(1)
r = edu.add_run("Bachelor of Technology in Mechanical Engineering")
r.bold = True; r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
r2 = edu.add_run("  |  J.N.T.U.H")
r2.font.size = Pt(10.5); r2.font.name = "Calibri"; r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

edu2 = doc.add_paragraph()
edu2.paragraph_format.space_after = Pt(0)
r = edu2.add_run("Christu Jyothi Institute of Technology and Science (CJITS), Jangaon  |  Graduated 2018")
r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ═══════════════════════════════════════════════════════
# AWARDS & ACCOMPLISHMENTS
# ═══════════════════════════════════════════════════════
section_heading("Awards & Accomplishments")
awards = [
    "Promoted to Quality Check specialist due to excellent quality and productivity.",
    "Received financial reward and recognition from upper management for leading a team to achieve exceptional results.",
    "Recognized as Employee of the Month for outstanding performance and contributions to team goals.",
    "Maintained consistently high customer satisfaction ratings throughout my tenure.",
]
for a in awards:
    add_bullet(a)

# ═══════════════════════════════════════════════════════
# ADDITIONAL INFORMATION
# ═══════════════════════════════════════════════════════
section_heading("Additional Information")
add_skill_line("Languages", "English, Hindi, Telugu")
add_skill_line("Hobbies", "Playing sports, Traveling, Listening to music")

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"Resume saved to: {OUTPUT_PATH}")
